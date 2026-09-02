# -*- coding: utf-8 -*-
"""八十吉 パフォーマー指名予約 ご確認・ご記入シート（外部ツールなし版）"""
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = 'メイリオ'
NAVY = RGBColor(0x1F, 0x2A, 0x5C)
ACCENT = RGBColor(0xBF, 0x4E, 0x0E)
MUTED = RGBColor(0x55, 0x60, 0x7A)
RULE = RGBColor(0x88, 0x90, 0xA6)

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)
sec.top_margin = sec.bottom_margin = Inches(0.7)
sec.left_margin = sec.right_margin = Inches(0.75)

st = doc.styles['Normal']
st.font.name = FONT
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)


def _fmt(run, size=11, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    return run


def _p(indent=0.0, before=2, after=2):
    pg = doc.add_paragraph()
    pf = pg.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    pf.line_spacing = 1.15
    if indent:
        pf.left_indent = Inches(indent)
    return pg


def _shade(pg, fill):
    pPr = pg._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def para(text='', size=11, bold=False, color=None, before=2, after=2, indent=0.0):
    pg = _p(indent, before, after)
    if text:
        _fmt(pg.add_run(text), size, bold, color)
    return pg


def line(*parts, indent=0.0, size=11, before=2, after=2):
    pg = _p(indent, before, after)
    for value, kind in parts:
        if kind == 'rule':
            _fmt(pg.add_run('＿' * value), size, False, RULE)
        else:
            _fmt(pg.add_run(value), size)
    return pg


def blank(label, width=26, size=11, indent=0.0):
    return line((label, 'text'), (width, 'rule'), indent=indent, size=size)


def check(options, indent=0.22, size=10.5, before=2, after=2):
    pg = _p(indent, before, after)
    _fmt(pg.add_run('　'.join('□ ' + o for o in options)), size)
    return pg


def checks(options, indent=0.22, size=10.5):
    for i, o in enumerate(options):
        pg = _p(indent, 1 if i else 3, 1)
        _fmt(pg.add_run('□ ' + o), size)


def chapter(text):
    pg = _p(0, 16, 8)
    _fmt(pg.add_run('  ' + text), 13, True, RGBColor(0xFF, 0xFF, 0xFF))
    _shade(pg, '1F2A5C')
    return pg


def section(num, title):
    pg = _p(0, 12, 5)
    _fmt(pg.add_run('  %s. %s' % (num, title)), 12, True, RGBColor(0xFF, 0xFF, 0xFF))
    _shade(pg, '55607A')
    return pg


_seq = [0]


def item(title):
    _seq[0] += 1
    pg = _p(0, 12, 3)
    _fmt(pg.add_run('  【%d】%s' % (_seq[0], title)), 11.5, True, NAVY)
    _shade(pg, 'EEF1F7')
    return pg


def confirm():
    pg = _p(0.15, 5, 1)
    _fmt(pg.add_run('□ この内容で問題ありません'), 10.5)
    pg = _p(0.15, 1, 4)
    _fmt(pg.add_run('□ 気になる点があります（'), 10.5)
    _fmt(pg.add_run('＿' * 17), 10.5, color=RULE)
    _fmt(pg.add_run('）'), 10.5)


def note(text, size=9.5, indent=0.15):
    return para(text, size, False, MUTED, before=0, after=4, indent=indent)


def keep(*pgs):
    for pg in pgs:
        pg._p.get_or_add_pPr().append(OxmlElement('w:keepNext'))


# =====================================================================
para('八十吉やそきち　LINE公式アカウント構築', 12, False, MUTED, 0, 0)
para('ご確認・ご記入シート', 20, True, NAVY, 0, 6)
para('パフォーマーのご指名予約をLINEで回すにあたり、'
     'ご確認いただきたい内容とお教えいただきたい情報をまとめました。',
     10.5, False, MUTED, 0, 2)
para('はじめの3件を打ち合わせでうかがい、第1部はチェック、第2部はご記入をお願いいたします。'
     'すべて埋まらなくても構いませんので、分かる範囲でお願いします。',
     10.5, False, MUTED, 0, 10)

line(('ご記入日：', 'text'), (3, 'rule'), ('年', 'text'), (2, 'rule'), ('月', 'text'),
     (2, 'rule'), ('日', 'text'), ('　　ご記入者：', 'text'), (12, 'rule'), after=6)

# =====================================================================
chapter('はじめに、特にうかがいたいこと（3件）')
para('この3つが決まらないと、仕組みの組み立てに入れません。'
     '打ち合わせの場でうかがえますと助かります。', 10.5, False, MUTED, 0, 4)

k = [item('ご予約サイトで「パフォーマーのご指名」をどう受けますか')]
k.append(para('ご予約はこれまで通り食べログ・ホットペッパー等で承ります。'
              'ただし、ご予約サイトには「担当者を指名する」機能が通常ございません。'
              'どのように指名をお受けになるかで、その後の流れが変わります。', 10.5))
keep(*k)
checks(['ご要望・備考欄に書いていただく',
        'コースとして登録する（「〇〇さん指名プラン」など）',
        'まだ決めていない　／　ご相談したい'], indent=0.28)

k = [item('いまお使いのご予約サイトを教えてください')]
keep(*k)
check(['食べログ', 'ホットペッパーグルメ', 'ぐるなび'], indent=0.28)
line(('□ その他（', 'text'), (18, 'rule'), ('）', 'text'), indent=0.28, size=10.5)
para('複数のサイトの空席をまとめて管理する仕組み（サイトコントローラー）はお使いですか。',
     10.5, before=6)
check(['使っている', '使っていない', '分からない'], indent=0.28)
line(('　お使いの場合、名称：', 'text'), (20, 'rule'), indent=0.28, size=10.5)
note('お席が二重に埋まらないようにするため、ここを確認させてください。', indent=0.28)

k = [item('パフォーマーの方のお名前を、LINEでご案内してもよろしいですか')]
k.append(para('掲載できると「この日に◯◯さんが出ます」というご案内ができ、'
              'ご指名につながります。掲載できない場合は'
              '「イベントがあります」という案内にとどまります。', 10.5))
keep(*k)
checks(['実名・芸名で掲載してよい', 'ご本人に確認のうえ、可の方のみ掲載したい',
        '掲載しない方がよい'], indent=0.28)

# =====================================================================
doc.add_page_break()
chapter('第1部　構築の前提について（チェックをお願いします）')
para('弊社側で決めている内容です。認識に相違がないかご確認ください。', 10.5, False, MUTED, 0, 4)

k = [item('パフォーマーの方にLINEの操作権限をお渡しします')]
k.append(para('ご予約が入った際に、パフォーマーの方ご自身がLINEでお客様とやり取り'
              'できるようにいたします。個人のLINEを教える必要がなく、'
              'やり取りは店舗のアカウントに残るため、ほかの従業員の方もご確認いただけます。', 10.5))
k.append(para('ただし、権限をお持ちの方は、ほかのお客様とのやり取りもご覧いただける状態に'
              'なります。またお辞めになった際には、権限を外す作業が必要です。'
              'この2点をご了承いただけますでしょうか。', 10.5, before=5))
keep(*k)
confirm()

k = [item('ご予約はこれまで通り、ご予約サイトで承ります')]
k.append(para('LINEではご予約をお受けしません。'
              'お席の在庫はご予約サイトが持ったままですので、'
              'お席が二重に埋まることがありません。', 10.5))
keep(*k)
confirm()

k = [item('LINEが担うのは「連絡」です')]
k.append(para('ご予約が入ってから、パフォーマーの方へのご連絡、'
              'ご承諾のお返事、店長様へのお知らせ。この3つをLINEで行います。'
              'ご予約の受付とお席の管理は、いまの仕組みのままです。', 10.5))
keep(*k)
confirm()

k = [item('外部のツールは導入せず、月額の固定費は発生しません')]
k.append(para('予約管理ツールなどの外部サービスは使いません。'
              'LINE公式アカウントは月額0円のプランから始められます。'
              'お客様が増えて一斉配信を始める段階で、はじめて費用が発生します。', 10.5))
keep(*k)
confirm()

k = [item('エイトビート様は一旦外し、八十吉様のみで構築します')]
k.append(para('店舗の開設時期が定まってからの対応といたします。'
              '仕組みは同じものを使えますので、開設後に移す形になります。', 10.5))
keep(*k)
confirm()

k = [item('今回は構築までで、納品後の配信代行は含みません')]
k.append(para('構築と運用マニュアルのお渡しまでが範囲です。'
              '納品後の運用は店舗様にてご対応いただきます。', 10.5))
keep(*k)
confirm()

k = [item('一斉配信は、新規のお客様に向けたものだけといたします')]
k.append(para('再来店をうながす配信やステップ配信は行いません。'
              'パフォーマーの募集と、新しいお客様への告知に絞ります。', 10.5))
keep(*k)
confirm()

# =====================================================================
doc.add_page_break()
chapter('第2部　お教えいただきたい情報（ご記入をお願いします）')

section(1, '八十吉様の基本情報')
para('媒体によって記載が異なっており、特に日曜日が営業日か定休日かが分かれております。',
     10, False, MUTED, after=4)
blank('営業時間　　　　　　：')
blank('定休日　　　　　　　：')
line(('お席　　　　　　　　：', 'text'), (4, 'rule'), ('席　（1組あたり　', 'text'),
     (3, 'rule'), ('名くらい）', 'text'))

section(2, 'パフォーマーについて')
line(('何名を想定されていますか　　　　：', 'text'), (4, 'rule'), ('名', 'text'))
para('出勤できる日は、どのように把握されますか', before=8, after=2)
checks(['月ごとに提出してもらう', 'その都度こちらから聞く', 'まだ決めていない'], indent=0.15)
line(('お支払いの目安（時給など）　　　：', 'text'), (14, 'rule'), before=8)
blank('募集で求めたいパフォーマンス　　：', 22, indent=0.0)
note('例：マジック、弾き語り、落語　など', indent=0.2)
para('※ お名前の掲載可否は、1ページ目の【3】でうかがっています', 9.5, False, MUTED,
     before=6, after=4)

section(3, 'キャンセルの取り決め')
para('お客様側', bold=True, before=4)
line(('いつまでのキャンセルなら無料にしますか　：', 'text'), (10, 'rule'), indent=0.15)
para('それを過ぎた場合', indent=0.15, before=5, after=2)
check(['キャンセル料をいただく（　　　　円 ／ 　　％）', '特にいただかない', '未定'],
      indent=0.32)
para('ご予約時のお支払い（カード払いなど）', indent=0.15, before=6, after=2)
check(['導入したい', '導入しない', '未定'], indent=0.32)
note('ご予約サイト側にお支払いの機能があるかによって、できることが変わります。', indent=0.32)

para('パフォーマー側', bold=True, before=10)
para('パフォーマーの方が直前に来られなくなった場合', indent=0.15, after=2)
checks(['報酬から差し引くなどの取り決めをしたい',
        '取り決めはせず、その都度対応する',
        'まだ決めていない'], indent=0.32)

section(4, 'その他')
para('LINEでこれもできたら、というご要望があればご記入ください')
for _ in range(3):
    blank('', 41)

para('ご確認・ご記入ありがとうございました。', 10.5, False, MUTED, before=14)
para('株式会社DYM', 10.5, False, MUTED, before=0)

out = sys.argv[1]
doc.save(out)
print('SAVED:', out)
