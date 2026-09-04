# -*- coding: utf-8 -*-
"""八十吉 LINE公式アカウント ご確認シート（簡潔版）"""
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = 'メイリオ'
NAVY = RGBColor(0x1F, 0x2A, 0x5C)
ACCENT = RGBColor(0xBF, 0x4E, 0x0E)
MUTED = RGBColor(0x55, 0x60, 0x7A)
RULE = RGBColor(0x88, 0x90, 0xA6)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)
sec.top_margin = sec.bottom_margin = Inches(0.75)
sec.left_margin = sec.right_margin = Inches(0.8)

st = doc.styles['Normal']
st.font.name = FONT
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

_seq = [0]


def _fmt(run, size=11, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    return run


def _p(indent=0.0, before=2, after=2):
    pg = doc.add_paragraph()
    pf = pg.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    pf.line_spacing = 1.2
    if indent:
        pf.left_indent = Inches(indent)
    return pg


def _shade(pg, fill):
    pPr = pg._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def _keep(*pgs):
    for pg in pgs:
        pg._p.get_or_add_pPr().append(OxmlElement('w:keepNext'))


def para(text='', size=11, bold=False, color=None, before=2, after=2, indent=0.0):
    pg = _p(indent, before, after)
    if text:
        _fmt(pg.add_run(text), size, bold, color)
    return pg


def line(*parts, indent=0.0, size=11, before=2, after=2):
    pg = _p(indent, before, after)
    for value, kind in parts:
        if kind == 'rule':
            _fmt(pg.add_run('＿' * value), size, False, RULE)
        else:
            _fmt(pg.add_run(value), size)
    return pg


def blank(label, width=26, indent=0.0):
    return line((label, 'text'), (width, 'rule'), indent=indent)


def chapter(text):
    pg = _p(0, 18, 8)
    _fmt(pg.add_run('  ' + text), 13, True, WHITE)
    _shade(pg, '1F2A5C')
    return pg


def item(title):
    _seq[0] += 1
    pg = _p(0, 14, 4)
    _fmt(pg.add_run('  【%d】%s' % (_seq[0], title)), 11.5, True, NAVY)
    _shade(pg, 'EEF1F7')
    return pg


def confirm():
    pg = _p(0.15, 6, 2)
    _fmt(pg.add_run('□ この内容で問題ありません'), 10.5)
    pg = _p(0.15, 2, 6)
    _fmt(pg.add_run('□ 気になる点があります（'), 10.5)
    _fmt(pg.add_run('＿' * 18), 10.5, color=RULE)
    _fmt(pg.add_run('）'), 10.5)


def bullet(text, size=11, indent=0.18):
    pg = _p(indent, 3, 3)
    _fmt(pg.add_run('・' + text), size)
    return pg


# =====================================================================
para('八十吉やそきち　LINE公式アカウント構築', 11.5, False, MUTED, 0, 0)
para('ご確認シート', 20, True, NAVY, 0, 6)
para('ご指名でのご予約をLINEでお受けするにあたり、'
     'ご確認いただきたい内容をまとめました。', 10.5, False, MUTED, 0, 10)

line(('ご記入日：', 'text'), (3, 'rule'), ('年', 'text'), (2, 'rule'), ('月', 'text'),
     (2, 'rule'), ('日', 'text'), ('　　ご記入者：', 'text'), (12, 'rule'), after=4)

# =====================================================================
chapter('ご確認をお願いします')
para('弊社側で決めている内容です。認識に相違がないかご確認ください。',
     10.5, False, MUTED, 0, 4)

k = [item('パフォーマーの方にLINEの操作権限をお渡しします')]
k.append(para('ご予約が入った際に、パフォーマーの方ご自身がLINEでお客様とやり取り'
              'できるようにいたします。個人のLINEを教える必要がなく、'
              'やり取りは店舗のアカウントに残るため、ほかの従業員の方もご確認いただけます。',
              10.5))
k.append(para('ただし、権限をお持ちの方は、ほかのお客様とのやり取りもご覧いただける'
              '状態になります。またお辞めになった際には、権限を外す作業が必要です。'
              'この2点をご了承いただけますでしょうか。', 10.5, before=4))
_keep(*k)
confirm()

k = [item('ご予約はこれまで通りで、LINEは連絡だけを担います')]
k.append(para('LINEの中に、ご予約の仕組みは作りません。'
              'ご希望はトークでお受けしますが、'
              'お席の確保はこれまで通りの方法で行っていただきます。', 10.5))
k.append(para('LINEが担うのは、ご指名の連絡と、'
              'パフォーマーの方とお客様のやり取りです。'
              'ご承諾のお返事も、ご来店時刻の調整も、確定のご連絡も、'
              'すべてトークの中で行われます。', 10.5, before=4))
_keep(*k)
confirm()

k = [item('外部のツールは導入せず、月額の固定費は発生しません')]
k.append(para('予約管理などの外部サービスは使いません。'
              'LINE公式アカウントは月額0円のプランから始められます。', 10.5))
k.append(para('費用が発生するのは、お客様全員へまとめてお送りする配信が'
              '月200通を超えたときです。'
              'お友だちが156名の場合、1回の配信で156通となります。',
              10.5, before=4))
k.append(para('通数を使うのは、一斉配信・絞り込み配信・ステップ配信の3つです。'
              'あいさつメッセージ、自動応答、お客様との個別のやり取りは、'
              '何通でも費用がかかりません。', 10.5, before=4))
_keep(*k)
confirm()

k = [item('一斉配信は、新規のお客様に向けたものだけといたします')]
k.append(para('再来店をうながす配信や、ステップ配信は行いません。'
              '新しいお客様への告知と、出演者の募集に絞ります。'
              'なお、ご応募の受付はLINEでは行いません。', 10.5))
_keep(*k)
confirm()

# =====================================================================
chapter('お教えいただきたいこと')

k = [item('八十吉様の営業時間と定休日')]
k.append(para('ご予約サイトによって記載が異なっており、'
              '特に日曜日が営業日か定休日かが分かれております。'
              '正しい情報をお教えください。', 10.5))
_keep(*k)
blank('営業時間　：', 24, indent=0.15)
blank('定休日　　：', 24, indent=0.15)

# =====================================================================
chapter('エイトビート様を始められるときに決めること')
para('出演者の方や内容が決まってから、あらためてご相談させてください。'
     'いまお決めいただく必要はありません。', 10.5, False, MUTED, 0, 5)

bullet('パフォーマーの人数と、出勤の決め方')
bullet('ご指名の書き方（お名前のみ／芸名／番号など）')
bullet('キャンセルの取り決め（お客様側・パフォーマー側）')

para('これらが決まりましたら、文面とメニューの表記を差し替えるだけで運用に入れます。',
     10.5, before=6)

para('', before=12)
para('ご確認ありがとうございました。', 10.5, False, MUTED)
para('株式会社DYM', 10.5, False, MUTED, before=2)

out = sys.argv[1]
doc.save(out)
print('SAVED:', out)
