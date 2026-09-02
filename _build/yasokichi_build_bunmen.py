# -*- coding: utf-8 -*-
"""八十吉 パフォーマー指名予約 文面集"""
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = 'メイリオ'
NAVY = RGBColor(0x1F, 0x2A, 0x5C)
ACCENT = RGBColor(0xBF, 0x4E, 0x0E)
MUTED = RGBColor(0x55, 0x60, 0x7A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)
sec.top_margin = sec.bottom_margin = Inches(0.7)
sec.left_margin = sec.right_margin = Inches(0.75)

st = doc.styles['Normal']
st.font.name = FONT
st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

_no = [0]


def _fmt(run, size=10.5, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    return run


def _p(indent=0.0, before=2, after=2, spacing=1.25):
    pg = doc.add_paragraph()
    pf = pg.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    pf.line_spacing = spacing
    if indent:
        pf.left_indent = Inches(indent)
    return pg


def _shade(pg, fill):
    pPr = pg._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def _bar(pg, color):
    """段落の左に縦罫を入れる"""
    pPr = pg._p.get_or_add_pPr()
    bd = OxmlElement('w:pBdr')
    lf = OxmlElement('w:left')
    lf.set(qn('w:val'), 'single'); lf.set(qn('w:sz'), '18')
    lf.set(qn('w:space'), '8'); lf.set(qn('w:color'), color)
    bd.append(lf)
    pPr.append(bd)


def para(text='', size=10.5, bold=False, color=None, before=2, after=2, indent=0.0):
    pg = _p(indent, before, after)
    if text:
        _fmt(pg.add_run(text), size, bold, color)
    return pg


def h1(text):
    pg = _p(0, 16, 7)
    _fmt(pg.add_run('  ' + text), 12.5, True, WHITE)
    _shade(pg, '1F2A5C')
    return pg


def item(title, note=None):
    """文面の見出し（自動採番）"""
    _no[0] += 1
    pg = _p(0, 13, 3)
    _fmt(pg.add_run('  %02d.　%s' % (_no[0], title)), 11.5, True, NAVY)
    _shade(pg, 'EEF1F7')
    if note:
        para(note, 9.5, False, MUTED, before=1, after=2, indent=0.05)
    return pg


def msg(text, star=False):
    """送信する文面そのもの。左に縦罫を入れて本文と区別する"""
    lines = text.split('\n')
    for i, ln in enumerate(lines):
        pg = _p(0.12, 1 if i else 4, 4 if i == len(lines) - 1 else 1, spacing=1.2)
        _bar(pg, 'BF4E0E' if star else 'A9B2C6')
        _fmt(pg.add_run(ln if ln else '　'), 10.5)
    return


def note(text):
    return para('※ ' + text, 9.5, False, MUTED, before=1, after=4, indent=0.05)


def star(text):
    return para('★ ' + text, 9.5, False, ACCENT, before=1, after=4, indent=0.05)


# =====================================================================
para('八十吉やそきち　LINE公式アカウント', 11, False, MUTED, 0, 0)
para('パフォーマー指名予約　文面集', 18, True, NAVY, 0, 4)
para('2026年9月2日　／　株式会社DYM　WEB事業部 SNS課', 10, False, MUTED, 0, 8)

para('LINE公式アカウントに登録する文面を、送信する場面ごとにまとめました。'
     'そのままお使いいただける形にしてありますので、'
     '言い回しなど気になる箇所に直接お書き込みください。')
para('★の付いた箇所は、店舗様のご確認をいただかないと確定できない部分です。',
     bold=True, color=ACCENT, before=4)

para('本文面は、お客様からのご予約リクエストをLINEでお受けし、'
     '店舗様がご予約媒体へご入力いただく流れを前提としています。'
     'ご予約媒体で直接お受けになる場合は、03の文面のみ差し替えとなります。',
     10, False, MUTED, before=5)

# =====================================================================
h1('第1部　お客様向け（自動で送られるもの）')

item('あいさつメッセージ　1通目',
     '友だち追加された直後に自動送信されます。')
msg('【お名前】さん\n\n'
    '友だち追加ありがとうございます。\n'
    'あさぶの居酒屋 八十吉です。')
note('【お名前】の部分には、お客様のLINEの表示名が自動で入ります。')

item('あいさつメッセージ　2通目',
     '1通目に続けて自動送信されます。')
msg('このアカウントでは、ご予約の受付と\n'
    '出演者の情報をお届けしております。\n\n'
    'ご予約は下のメニューの「ご予約」から\n'
    '24時間承っております。')

item('ご予約（メニューを押したとき）',
     'リッチメニューの「ご予約」を押すと自動で表示されます。ご予約の入口となる、'
     'もっとも重要な文面です。')
msg('ご予約ありがとうございます。\n\n'
    '下記をコピーして、埋めたうえでお送りください。\n\n'
    '－－－－－－－－－－\n'
    'ご希望日：　　月　　日（　）\n'
    'お時間　：　　時　　分ごろ\n'
    '人数　　：　　名\n'
    'ご指名　：\n'
    'ご要望　：\n'
    '－－－－－－－－－－\n\n'
    '内容を確認のうえ、担当者よりお返事いたします。\n\n'
    '※お席と出演者の空き状況を確認いたしますので、\n'
    '　この時点ではご予約は確定しておりません。')
note('最後の2行が要点です。リクエスト制であることをここで明示しておくと、'
     'あとで「予約したつもりだった」という行き違いを防げます。')

item('出演者のご紹介（メニューを押したとき）')
msg('【出演者のご紹介】\n\n'
    '★準備中です。決まり次第こちらでご案内いたします。', star=True)
star('出演者が決まり次第、お名前・ご経歴・演目を差し替えます。'
     'お名前を掲載してよいかは、ご確認シートの【3】でうかがっています。')

item('はじめての方へ（メニューを押したとき）')
msg('【あさぶの居酒屋 八十吉】\n\n'
    '◇ 場所\n'
    '札幌市北区麻生5-7-3 5丁目ビル1F\n'
    '地下鉄南北線 麻生駅1番出口から徒歩1分\n\n'
    '◇ 営業時間\n'
    '★確認中\n\n'
    '◇ ご予算\n'
    'お一人さま 4,000円〜\n\n'
    'ご予約は下のメニューの「ご予約」から承っております。', star=True)
star('営業時間は、ご予約媒体によって記載が異なっておりました。'
     '特に日曜日が営業日か定休日かが分かれております。正しい情報をお教えください。')

item('お問い合わせ（メニューを押したとき）')
msg('ご質問・ご相談は、このままメッセージを\n'
    'お送りください。\n\n'
    '営業時間内に担当者よりお返事いたします。')

item('営業時間外の自動返信',
     '営業時間外にメッセージが届いたとき、自動で返します。'
     'この返信は通数としてカウントされないため、費用は発生しません。')
msg('メッセージをありがとうございます。\n\n'
    'ただいま営業時間外のため、翌営業日に\n'
    '順次お返事いたします。\n\n'
    'ご予約のご希望は、下のメニューの「ご予約」から\n'
    'いつでもお送りいただけます。')

item('「予約」と送られたときの自動応答',
     'メニューを使わずに「予約」「よやく」などと入力された場合に返します。')
msg('ご予約をご希望でしょうか。\n\n'
    '下のメニューの「ご予約」を押していただくと、\n'
    'ご記入いただく項目をご案内いたします。')

item('「場所」「アクセス」と送られたときの自動応答')
msg('【アクセス】\n\n'
    '札幌市北区麻生5-7-3 5丁目ビル1F\n'
    '地下鉄南北線 麻生駅1番出口から徒歩1分です。')

# =====================================================================
doc.add_page_break()
h1('第2部　出演者の方向け（店舗様から送るもの）')

para('出演者の方には、店舗のアカウントの操作権限をお渡しします。'
     'そのうえで、下記を個別のトークでお送りいただきます。', 10, False, MUTED, after=4)

item('ご指名が入ったことのお知らせ',
     'ご予約のリクエストを受け取ったら、該当の出演者へお送りします。')
msg('【ご指名が入りました】\n\n'
    '◇ 日時\n'
    '　　月　　日（　）　　時　　分〜\n\n'
    '◇ 人数\n'
    '　　　名\n\n'
    '◇ ご要望\n'
    '　\n\n'
    'お受けいただけますでしょうか。\n'
    '★本日　　時までにご返信をお願いいたします。', star=True)
star('返信の期限を何時にするかを決めてください。'
     'お客様をお待たせする時間に直結します。')

item('お返事がないときの確認',
     '期限が近づいても返信がない場合にお送りします。')
msg('先ほどのご指名の件、ご確認いただけましたでしょうか。\n\n'
    'お客様をお待たせしておりますので、\n'
    '難しい場合もその旨お知らせいただけますと助かります。')

item('ご予約が確定したことのお知らせ')
msg('【確定しました】\n\n'
    '◇ 日時\n'
    '　　月　　日（　）　　時　　分〜\n\n'
    '◇ 人数\n'
    '　　　名\n\n'
    'よろしくお願いいたします。\n'
    '当日の入り時間などは、追ってご連絡いたします。')

item('お客様都合でキャンセルになったときのお知らせ')
msg('【キャンセルのご連絡】\n\n'
    '　　月　　日（　）のご予約が、\n'
    'お客様のご都合によりキャンセルとなりました。\n\n'
    '直前のご連絡となり申し訳ございません。')

# =====================================================================
h1('第3部　お客様への個別のご連絡（店舗様から送るもの）')

item('リクエストを受け取ったことのお知らせ',
     'お客様からご予約のリクエストが届いたら、まずこれをお送りします。')
msg('ご予約のリクエストをありがとうございます。\n\n'
    'お席と出演者の空き状況を確認のうえ、\n'
    '★本日　　時までにお返事いたします。\n\n'
    'しばらくお待ちくださいませ。', star=True)
star('何時までに返事をするかを決めてください。'
     'ここを明示できると、お客様の不安がなくなります。')

item('ご予約確定のお知らせ')
msg('お待たせいたしました。\n'
    '下記の内容でご予約を承りました。\n\n'
    '◇ ご予約日時\n'
    '　　月　　日（　）　　時　　分\n\n'
    '◇ 人数\n'
    '　　　名\n\n'
    '◇ ご指名\n'
    '　\n\n'
    'ご来店を心よりお待ちしております。\n'
    '麻生駅1番出口から徒歩1分です。\n\n'
    '※ご変更・キャンセルは、このトークにて\n'
    '　お知らせください。')

item('満席でお受けできないときのお知らせ')
msg('ご予約のリクエストをありがとうございました。\n\n'
    '大変申し訳ございませんが、ご希望のお日にちは\n'
    '満席となっております。\n\n'
    '別のお日にちでしたらご用意できる場合がございます。\n'
    'ご希望をお聞かせいただけますと幸いです。')

item('出演者の都合でお受けできないときのお知らせ')
msg('ご予約のリクエストをありがとうございました。\n\n'
    '大変申し訳ございませんが、ご指名いただいた出演者は\n'
    'そのお日にちの出演がかなわない状況です。\n\n'
    'お席のご用意はできますので、\n'
    '別の出演者、または別のお日にちでご検討いただけますでしょうか。')
note('お席は空いている場合が多いので、そのまま来店につなげられる書き方にしています。')

item('前日のお知らせ',
     'ご来店の前日にお送りします。')
msg('明日のご予約が近づいてまいりました。\n\n'
    '◇ ご予約日時\n'
    '　　月　　日（　）　　時　　分\n\n'
    '◇ 人数\n'
    '　　　名\n\n'
    '◇ ご指名\n'
    '　\n\n'
    'お気をつけてお越しくださいませ。\n'
    'ご来店を心よりお待ちしております。')

item('キャンセルを承ったときのお知らせ')
msg('承知いたしました。\n\n'
    '　　月　　日（　）のご予約を\n'
    'キャンセルさせていただきました。\n\n'
    '★キャンセル料について\n\n'
    'またのご来店をお待ちしております。', star=True)
star('キャンセル料をいただくかどうか、いつまでなら無料かを決めてください。'
     'ご確認シートの第2部でうかがっています。')

# =====================================================================
h1('第4部　店舗内の共有')

item('ご予約が成立したときのお知らせ',
     '店長様など、店舗内で共有する方へお送りします。')
msg('【ご予約が成立しました】\n\n'
    '　　月　　日（　）　　時　　分／　　名\n'
    '　ご指名：\n'
    '　ご要望：\n\n'
    'ご予約媒体への入力：済')

# =====================================================================
doc.add_page_break()
h1('★　ご確認をお願いしたい箇所')

para('文面を確定するために、下記をお教えください。', after=5)

for t in [
    '04　出演者のお名前・ご経歴・演目（掲載の可否はご確認シート【3】でうかがっています）',
    '05　八十吉様の正しい営業時間と定休日（ご予約媒体によって記載が異なっておりました）',
    '10　出演者の方に、何時までにご返信いただくか',
    '13　お客様に、何時までにお返事するか',
    '18　キャンセル料をいただくか。いつまでなら無料とするか',
]:
    pg = _p(0.12, 3, 3)
    _bar(pg, 'BF4E0E')
    _fmt(pg.add_run(t), 10.5)

para('', before=8)
para('このほか、言い回しでお気づきの点がございましたら、'
     '該当箇所に直接お書き込みください。', 10, False, MUTED)

para('', before=10)
para('以上', 10.5, False, MUTED)
para('株式会社DYM　WEB事業部 SNS課', 10, False, MUTED, before=4)

out = sys.argv[1]
doc.save(out)
print('SAVED:', out, '／ 文面', _no[0], '本')
