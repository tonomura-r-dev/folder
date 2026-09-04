# -*- coding: utf-8 -*-
"""八十吉 パフォーマーへの権限付与に関するご確認（1枚）"""
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = 'メイリオ'
NAVY = RGBColor(0x1F, 0x2A, 0x5C)
MUTED = RGBColor(0x55, 0x60, 0x7A)
RULE = RGBColor(0x88, 0x90, 0xA6)

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)
sec.top_margin = sec.bottom_margin = Inches(0.9)
sec.left_margin = sec.right_margin = Inches(0.85)

st = doc.styles['Normal']
st.font.name = FONT
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)


def _fmt(run, size=11, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    return run


def _p(indent=0.0, before=2, after=2):
    pg = doc.add_paragraph()
    pf = pg.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    pf.line_spacing = 1.3
    if indent:
        pf.left_indent = Inches(indent)
    return pg


def _shade(pg, fill):
    pPr = pg._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def para(text='', size=11, bold=False, color=None, before=2, after=2, indent=0.0):
    pg = _p(indent, before, after)
    if text:
        _fmt(pg.add_run(text), size, bold, color)
    return pg


# =====================================================================
para('八十吉やそきち　LINE公式アカウント構築', 11.5, False, MUTED, 0, 0)
para('ご確認事項', 20, True, NAVY, 0, 14)

pg = _p(0, 10, 8)
_fmt(pg.add_run('  【1】パフォーマーの方にLINEの操作権限をお渡しします'), 13, True, NAVY)
_shade(pg, 'EEF1F7')

para('ご予約が入った際に、パフォーマーの方ご自身がLINEでお客様とやり取りできる'
     'ようにいたします。個人のLINEを教える必要がなく、やり取りは店舗のアカウントに'
     '残るため、ほかの従業員の方もご確認いただけます。', 11.5, before=6, after=8)

para('ただし、権限をお持ちの方は、ほかのお客様とのやり取りもご覧いただける状態に'
     'なります。またお辞めになった際には、権限を外す作業が必要です。', 11.5, after=8)

para('この2点をご了承いただけますでしょうか。', 11.5, True, NAVY, after=12)

pg = _p(0.2, 6, 4)
_fmt(pg.add_run('□　この内容で問題ありません'), 11.5)
pg = _p(0.2, 4, 4)
_fmt(pg.add_run('□　気になる点があります（'), 11.5)
_fmt(pg.add_run('＿' * 20), 11.5, color=RULE)
_fmt(pg.add_run('）'), 11.5)

para('', before=26)
pg = _p(0, 6, 2)
for txt, kind in [('ご記入日：', 't'), (4, 'r'), ('年', 't'), (3, 'r'), ('月', 't'),
                  (3, 'r'), ('日', 't'), ('　　　ご記入者：', 't'), (14, 'r')]:
    if kind == 'r':
        _fmt(pg.add_run('＿' * txt), 11.5, color=RULE)
    else:
        _fmt(pg.add_run(txt), 11.5)

para('株式会社DYM', 10.5, False, MUTED, before=30)

out = sys.argv[1]
doc.save(out)
print('SAVED:', out)
